"""
Document Extractor service.

Extracts text from PDF, DOCX, JPG, and PNG files using:
- PyMuPDF (fitz) for text-based PDFs
- pdf2image + pytesseract for scanned/image PDFs
- python-docx for DOCX files
- pytesseract directly for image files (JPG, PNG)

Usage:
    from app.services.document_extractor import extract
    text, confidence = extract(Path("/tmp/resume.pdf"), ".pdf")
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Lazy imports — keep top-level import errors readable
# ---------------------------------------------------------------------------
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore[assignment]

try:
    import docx  # python-docx
except ImportError:
    docx = None  # type: ignore[assignment]

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None  # type: ignore[assignment]

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
SUPPORTED_TYPES = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}
OCR_LOW_QUALITY_THRESHOLD = 60.0
OCR_LANG = "vie+eng"

# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract(file_path: Path, file_type: str) -> tuple[str, float]:
    """
    Extract text from a document file.

    Args:
        file_path: Path to the file on disk.
        file_type: File extension including dot, e.g. ".pdf", ".docx", ".jpg".

    Returns:
        A (text, confidence) tuple where:
        - text: Extracted plain text string.
        - confidence: Float 0–100. 1.0 (normalised to 100.0) for text-based
          extraction; mean OCR confidence for image-based extraction.

    Raises:
        ValueError: If file_type is not in SUPPORTED_TYPES.
    """
    file_type = file_type.lower()

    if file_type not in SUPPORTED_TYPES:
        raise ValueError(
            f"Unsupported file type: {file_type!r}. "
            f"Supported types: {sorted(SUPPORTED_TYPES)}"
        )

    if file_type == ".pdf":
        return _extract_pdf(file_path)
    if file_type in (".docx", ".doc"):
        return _extract_docx(file_path)
    if file_type in (".jpg", ".jpeg", ".png"):
        return _extract_image(file_path)

    # Should never reach here given the check above, but for safety:
    raise ValueError(f"Unsupported file type: {file_type!r}")


# ---------------------------------------------------------------------------
# Internal extractors
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: Path) -> tuple[str, float]:
    """Extract text from a PDF file. Falls back to OCR for image-only PDFs."""
    doc = fitz.open(str(file_path))
    pages_text: list[str] = []

    try:
        for page in doc:
            pages_text.append(page.get_text())
    finally:
        doc.close()

    full_text = "\n".join(pages_text).strip()

    # If text is substantial, treat as text-based PDF
    if len(full_text) > 20:
        return full_text, 1.0

    # Otherwise fall back to OCR (scanned/image PDF)
    logger.info("PDF appears to be image-based, falling back to OCR: %s", file_path)
    return _extract_pdf_ocr(file_path)


def _extract_pdf_ocr(file_path: Path) -> tuple[str, float]:
    """OCR-based extraction for scanned/image PDFs using pdf2image + pytesseract."""
    images = convert_from_path(str(file_path))
    texts: list[str] = []
    confidences: list[float] = []

    for image in images:
        text, confidence = _ocr_image(image)
        texts.append(text)
        confidences.append(confidence)

    full_text = "\n".join(texts)
    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

    if avg_confidence < OCR_LOW_QUALITY_THRESHOLD:
        logger.warning(
            "OCR low quality (confidence=%.1f) for file: %s", avg_confidence, file_path
        )

    return full_text, avg_confidence


def _extract_docx(file_path: Path) -> tuple[str, float]:
    """Extract text from a DOCX/DOC file using python-docx."""
    document = docx.Document(str(file_path))
    parts: list[str] = []

    # Extract paragraph text
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract table cell text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    text = "\n".join(parts)
    return text, 1.0  # No OCR → full confidence


def _extract_image(file_path: Path) -> tuple[str, float]:
    """Extract text from an image file (JPG/PNG) using pytesseract."""
    image = Image.open(str(file_path))
    text, confidence = _ocr_image(image)

    if confidence < OCR_LOW_QUALITY_THRESHOLD:
        logger.warning(
            "OCR low quality (confidence=%.1f) for file: %s", confidence, file_path
        )

    return text, confidence


def _ocr_image(image: Any) -> tuple[str, float]:
    """Run pytesseract OCR on a PIL image and return (text, mean_confidence)."""
    text: str = pytesseract.image_to_string(image, lang=OCR_LANG)

    # Get confidence data; filter out empty strings and -1 (no valid word)
    data: dict = pytesseract.image_to_data(image, lang=OCR_LANG, output_type=pytesseract.Output.DICT)  # type: ignore[attr-defined]
    confs: list[float] = [
        float(c)
        for c, t in zip(data.get("conf", []), data.get("text", []))
        if str(c).strip() not in ("-1", "") and str(t).strip()
    ]
    avg_conf = sum(confs) / len(confs) if confs else 0.0

    return text, avg_conf
